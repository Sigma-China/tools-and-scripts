import os
import sys
import nltk
import time
from nltk.tag import pos_tag
from nltk.tokenize import PunktSentenceTokenizer, word_tokenize
from bs4 import BeautifulSoup
from docx import Document

sys.stdout.reconfigure(encoding="utf-8")

# 打开文档
file_path = r"C:\Users\huw\OneDrive - Sigma Technology\Desktop\Scripts\Hik-Connect Mobile Client_User Manual.docx"  # 替换成你的文件路径，例如 "C:/Users/User Manual.pdf" 或 r"C:/Users/User Manual.pdf"
print("文件路径：", file_path)

# 定义支持打开 .docx，.html，.txt 文件的函数
def extract_text_from_file(file_path):
    _, file_extension = os.path.splitext(file_path)

    if file_extension == ".txt":
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read()
    elif file_extension == ".html":
        with open(file_path, "r", encoding="utf-8") as file:
            soup = BeautifulSoup(file, "html.parser")
            return soup.get_text()
    elif file_extension == ".docx":
        doc = Document(file_path)
        text = " ".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    else:
        raise ValueError("Unsupported file type")

# 从文件中提取文本内容
text = extract_text_from_file(file_path)

# 获取当前时间戳
timestamp = ""

timestamp = time.strftime("%Y%m%d_%H%M%S", time.localtime())

# 将文本分割成句子，并使用PunktSentenceTokenizer计算句子数
tokenizer = PunktSentenceTokenizer()
sentences = tokenizer.tokenize(text)
num_sentences = len(sentences)

# 计算超过25个单词的句子数量
num_long_sentences = sum(1 for sentence in sentences if len(nltk.word_tokenize(sentence)) > 25)

sentence_score = num_long_sentences/num_sentences
sentence_score = round(sentence_score, 3)  # 保留小数点后3位
print("总句数:", num_sentences)
print("长句数:", num_long_sentences)
print(f"长句占比: {sentence_score}")
print(f"output-超过25词的长句-{timestamp}.docx")

# 创建一个新的Word文档
doc = Document()
# 初始化序号计数器
count = 1
# Print long sentences
for sentence in sentences:
    if len(nltk.word_tokenize(sentence)) > 25:
        result = f"\n{count}. 超过25词的长句: {sentence}"
        print(result)
        doc.add_paragraph(result)
        count += 1

# 保存Word文档，文件名中包含时间戳
doc.save(f"output-超过25词的长句-{timestamp}.docx")